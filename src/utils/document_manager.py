"""
Document Management System for FinSolve AI Assistant
Handles document upload, processing, chunking, embedding, and vector database updates.

Author: Dr. Erick K. Yegon
Version: 1.0.0
"""

import os
import shutil
import hashlib
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import pandas as pd
from loguru import logger

# Document processing imports
try:
    import PyPDF2
    import docx
    from openpyxl import load_workbook
except ImportError:
    logger.warning("Document processing libraries not available. Install PyPDF2, python-docx, openpyxl")

# Vector database imports
try:
    import chromadb
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document
    from langchain_community.vectorstores import Chroma
    from euriai.langchain_embed import EuriaiEmbeddings
except ImportError:
    logger.warning("Vector database libraries not available")


class DocumentManager:
    """Manages document upload, processing, and vector database updates"""
    
    def __init__(self, data_dir: str = "data", vector_db_dir: str = "vector_db"):
        self.data_dir = Path(data_dir)
        self.vector_db_dir = Path(vector_db_dir)
        self.supported_formats = {'.pdf', '.docx', '.txt', '.csv', '.xlsx', '.md'}
        
        # Initialize embeddings
        try:
            self.embeddings = EuriaiEmbeddings(api_key=os.getenv("EURI_API_KEY"))
        except Exception as e:
            logger.error(f"Failed to initialize embeddings: {e}")
            self.embeddings = None
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def upload_document(self, file_path: str, department: str, document_type: str = "policy") -> Dict[str, Any]:
        """
        Upload and process a new document
        
        Args:
            file_path: Path to the uploaded file
            department: Department the document belongs to
            document_type: Type of document (policy, report, data, etc.)
        
        Returns:
            Dictionary with upload status and metadata
        """
        try:
            file_path = Path(file_path)
            
            # Validate file format
            if file_path.suffix.lower() not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"Unsupported file format: {file_path.suffix}",
                    "supported_formats": list(self.supported_formats)
                }
            
            # Create department directory if it doesn't exist
            dept_dir = self.data_dir / department.lower()
            dept_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate unique filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            new_filename = f"{timestamp}_{file_path.name}"
            destination = dept_dir / new_filename
            
            # Copy file to data directory
            shutil.copy2(file_path, destination)
            
            # Extract text content
            content = self.extract_text(destination)
            
            if not content:
                return {
                    "success": False,
                    "error": "Failed to extract text from document"
                }
            
            # Create document metadata
            metadata = {
                "filename": new_filename,
                "original_name": file_path.name,
                "department": department,
                "document_type": document_type,
                "upload_date": datetime.now().isoformat(),
                "file_size": destination.stat().st_size,
                "file_hash": self.calculate_file_hash(destination),
                "content_length": len(content)
            }
            
            # Process and add to vector database
            success = self.add_to_vector_db(content, metadata)
            
            if success:
                # Log the upload
                logger.info(f"Document uploaded successfully: {new_filename} to {department}")
                
                return {
                    "success": True,
                    "filename": new_filename,
                    "metadata": metadata,
                    "message": f"Document '{file_path.name}' uploaded and indexed successfully"
                }
            else:
                # Remove file if vector DB update failed
                destination.unlink()
                return {
                    "success": False,
                    "error": "Failed to add document to vector database"
                }
                
        except Exception as e:
            logger.error(f"Document upload failed: {str(e)}")
            return {
                "success": False,
                "error": f"Upload failed: {str(e)}"
            }
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text content from various file formats"""
        try:
            suffix = file_path.suffix.lower()
            
            if suffix == '.txt' or suffix == '.md':
                return file_path.read_text(encoding='utf-8')
            
            elif suffix == '.pdf':
                return self.extract_pdf_text(file_path)
            
            elif suffix == '.docx':
                return self.extract_docx_text(file_path)
            
            elif suffix == '.csv':
                return self.extract_csv_text(file_path)
            
            elif suffix == '.xlsx':
                return self.extract_xlsx_text(file_path)
            
            else:
                logger.warning(f"Unsupported file format: {suffix}")
                return ""
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {str(e)}")
            return ""
    
    def extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF files"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            return ""
    
    def extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            return ""
    
    def extract_csv_text(self, file_path: Path) -> str:
        """Extract text from CSV files"""
        try:
            df = pd.read_csv(file_path)
            # Convert DataFrame to text representation
            text = f"CSV Data from {file_path.name}:\n\n"
            text += f"Columns: {', '.join(df.columns)}\n"
            text += f"Rows: {len(df)}\n\n"
            text += df.to_string(index=False)
            return text
        except Exception as e:
            logger.error(f"CSV extraction failed: {str(e)}")
            return ""
    
    def extract_xlsx_text(self, file_path: Path) -> str:
        """Extract text from XLSX files"""
        try:
            workbook = load_workbook(file_path, read_only=True)
            text = f"Excel Data from {file_path.name}:\n\n"
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                
                # Read data from sheet
                data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        data.append([str(cell) if cell is not None else "" for cell in row])
                
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0] if data else [])
                    text += df.to_string(index=False) + "\n\n"
            
            return text
        except Exception as e:
            logger.error(f"XLSX extraction failed: {str(e)}")
            return ""
    
    def calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            logger.error(f"Hash calculation failed: {str(e)}")
            return ""
    
    def add_to_vector_db(self, content: str, metadata: Dict[str, Any]) -> bool:
        """Add document content to vector database"""
        try:
            if not self.embeddings:
                logger.error("Embeddings not available")
                return False
            
            # Split content into chunks
            chunks = self.text_splitter.split_text(content)
            
            # Create Document objects
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = metadata.copy()
                doc_metadata.update({
                    "chunk_id": i,
                    "chunk_count": len(chunks),
                    "source": f"{metadata['department']}/{metadata['filename']}"
                })
                
                documents.append(Document(
                    page_content=chunk,
                    metadata=doc_metadata
                ))
            
            # Initialize or load vector store
            vector_store = Chroma(
                persist_directory=str(self.vector_db_dir),
                embedding_function=self.embeddings,
                collection_name="finsolve_documents"
            )
            
            # Add documents to vector store
            vector_store.add_documents(documents)
            
            logger.info(f"Added {len(documents)} chunks to vector database")
            return True
            
        except Exception as e:
            logger.error(f"Vector DB update failed: {str(e)}")
            return False
    
    def list_documents(self, department: Optional[str] = None) -> List[Dict[str, Any]]:
        """List all documents, optionally filtered by department"""
        try:
            documents = []
            
            if department:
                search_dirs = [self.data_dir / department.lower()]
            else:
                search_dirs = [d for d in self.data_dir.iterdir() if d.is_dir()]
            
            for dept_dir in search_dirs:
                if not dept_dir.exists():
                    continue
                    
                for file_path in dept_dir.iterdir():
                    if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                        stat = file_path.stat()
                        documents.append({
                            "filename": file_path.name,
                            "department": dept_dir.name,
                            "size": stat.st_size,
                            "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                            "path": str(file_path)
                        })
            
            return sorted(documents, key=lambda x: x['modified'], reverse=True)
            
        except Exception as e:
            logger.error(f"Document listing failed: {str(e)}")
            return []
    
    def delete_document(self, filename: str, department: str) -> Dict[str, Any]:
        """Delete a document and remove from vector database"""
        try:
            file_path = self.data_dir / department.lower() / filename
            
            if not file_path.exists():
                return {
                    "success": False,
                    "error": "Document not found"
                }
            
            # Remove file
            file_path.unlink()
            
            # TODO: Remove from vector database
            # This would require implementing document removal from ChromaDB
            
            logger.info(f"Document deleted: {filename} from {department}")
            
            return {
                "success": True,
                "message": f"Document '{filename}' deleted successfully"
            }
            
        except Exception as e:
            logger.error(f"Document deletion failed: {str(e)}")
            return {
                "success": False,
                "error": f"Deletion failed: {str(e)}"
            }


# Global document manager instance
document_manager = DocumentManager()
